import os
from typing import Dict, List, Optional
from docx import Document
import json
from pathlib import Path
from app.core.config import settings

class KnowledgeBase:
    def __init__(self):
        self.kb_dir = Path("data/Knowledge Base")
        # Store restaurant info, menu, and faqs associated with a restaurant key (e.g., "Barbeque Nation - New Delhi")
        self.restaurants: Dict[str, Dict] = {}
        self._load_knowledge_base()
    
    def _load_knowledge_base(self):
        """Load all knowledge base documents (restaurant info, menu, faqs) from DOCX files"""
        for file_path in self.kb_dir.glob("*.docx"):
            try:
                # Skip duplicate files (those with (1) in the name)
                if "(1)" in file_path.name:
                    continue

                doc = Document(file_path)
                filename = file_path.name
                restaurant_key = None

                # Determine the type of document and the associated restaurant key
                if "Barbeque Nation" in filename and ("New Delhi" in filename or "Bangalore" in filename):
                    # This is a main restaurant info file
                    restaurant_key = self._extract_restaurant_key_from_filename(filename)
                    if restaurant_key and restaurant_key not in self.restaurants:
                        self.restaurants[restaurant_key] = self._initialize_restaurant_data()
                    if restaurant_key:
                        # Process general info from this file and update the existing data
                        self._process_general_info(doc, self.restaurants[restaurant_key])

                elif "Menu" in filename:
                    # This is a menu file, associate with relevant cities
                    for city in settings.CITIES.keys():
                         if city.lower() in filename.lower() or ("Barbeque Nation" in filename and ("New Delhi" in filename or "Bangalore" in filename) is False):
                            # Associate general menu with all cities if filename is general, or with specific city
                            key = f"Barbeque Nation - {city.capitalize()}"
                            if key not in self.restaurants:
                                self.restaurants[key] = self._initialize_restaurant_data()
                            self._process_menu(doc, self.restaurants[key])
                            # If it's a general menu file, apply to all cities and break
                            if "Barbeque Nation" in filename and ("New Delhi" in filename or "Bangalore" in filename) is False:
                                 break # Assume general menu applies to all cities

                elif "FAQ" in filename:
                    # This is an FAQ file, associate with relevant cities
                    for city in settings.CITIES.keys():
                        if city.lower() in filename.lower() or ("Barbeque Nation" in filename and ("New Delhi" in filename or "Bangalore" in filename) is False):
                            # Associate general FAQ with all cities if filename is general, or with specific city
                            key = f"Barbeque Nation - {city.capitalize()}"
                            if key not in self.restaurants:
                                self.restaurants[key] = self._initialize_restaurant_data()
                            self._process_faqs(doc, self.restaurants[key])
                            # If it's a general FAQ file, apply to all cities and break
                            if "Barbeque Nation" in filename and ("New Delhi" in filename or "Bangalore" in filename) is False:
                                 break # Assume general FAQ applies to all cities

            except Exception as e:
                print(f"Error loading {file_path}: {str(e)}")
    
    def _initialize_restaurant_data(self) -> Dict:
        """Initialize the dictionary structure for a restaurant"""
        return {
            "name": "",
            "location": "",
            "address": "",
            "contact": "",
            "timings": "",
            "menu": [],
            "faqs": []
        }

    def _extract_restaurant_key_from_filename(self, filename: str) -> Optional[str]:
        """Extract a standardized restaurant key from a main restaurant info filename"""
        if "Barbeque Nation" in filename:
            if "New Delhi" in filename:
                return "Barbeque Nation - New Delhi"
            elif "Bangalore" in filename:
                return "Barbeque Nation - Bangalore"
        return None

    def _process_general_info(self, doc: Document, data: Dict):
        """Process general restaurant information from a DOCX document"""
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if "Location:" in text:
                data["location"] = text.replace("Location:", "").strip()
            elif "Address:" in text:
                data["address"] = text.replace("Address:", "").strip()
            elif "Contact:" in text:
                data["contact"] = text.replace("Contact:", "").strip()
            elif "Timings:" in text:
                data["timings"] = text.replace("Timings:", "").strip()

    def _process_menu(self, doc: Document, data: Dict):
        """Process menu information from a DOCX document"""
        current_section = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if "Menu:" in text:
                current_section = "menu"
            elif current_section == "menu" and text:
                data["menu"].append(text)
            # Stop processing menu if a new section header is found (simple approach)
            elif text.endswith(":") and current_section == "menu":
                 break

    def _process_faqs(self, doc: Document, data: Dict):
        """Process FAQ information from a DOCX document"""
        current_section = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if "FAQs:" in text:
                current_section = "faqs"
            elif current_section == "faqs" and text:
                data["faqs"].append(text)
            # Stop processing FAQs if a new section header is found (simple approach)
            elif text.endswith(":") and current_section == "faqs":
                 break

    def get_restaurant_info(self, restaurant_name: str) -> Optional[Dict]:
        """Get information about a specific restaurant (general info, menu, faqs)"""
        return self.restaurants.get(restaurant_name)
    
    def get_all_restaurants(self) -> List[str]:
        """Get list of all restaurant names (keys in the dictionary)"""
        return list(self.restaurants.keys())
    
    def get_restaurants_by_city(self, city: str) -> List[str]:
        """Get list of restaurant keys for a specific city"""
        return [name for name in self.restaurants.keys() if city.lower() in name.lower()]
    
    def get_menu(self, restaurant_name: str) -> List[str]:
        """Get menu items for a specific restaurant key"""
        restaurant = self.restaurants.get(restaurant_name)
        return restaurant.get("menu", []) if restaurant else []
    
    def get_faqs(self, restaurant_name: str) -> List[str]:
        """Get FAQs for a specific restaurant key"""
        restaurant = self.restaurants.get(restaurant_name)
        return restaurant.get("faqs", []) if restaurant else []

# Create a singleton instance
knowledge_base = KnowledgeBase() 